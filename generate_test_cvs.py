from docx import Document
from fpdf import FPDF
import os
from pathlib import Path

def create_docx(filename, content):
    doc = Document()
    doc.add_heading('Curriculum Vitae', 0)
    for section in content:
        doc.add_paragraph(section)
    doc.save(filename)

def create_pdf(filename, content):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    pdf.cell(200, 10, txt="Curriculum Vitae", ln=1, align='C')
    
    for section in content:
        pdf.multi_cell(0, 10, txt=section)
    pdf.output(filename)

cv_contents = {
    "pedreiro": [
        "Pedreiro Profissional - 10 anos de experiência",
        "Experiência Profissional:",
        "- Construção de alvenaria e acabamentos",
        "- Leitura e interpretação de plantas",
        "- Assentamento de pisos e revestimentos",
        "- Execução de obras residenciais e comerciais",
        "Formação: Curso Técnico em Construção Civil"
    ],
    "admin_geral": [
        "Profissional Multidisciplinar",
        "Experiência Profissional:",
        "- Coordenação de projetos diversos",
        "- Análise e documentação de processos",
        "- Experiência em gestão de equipes",
        "- Conhecimento em ferramentas administrativas",
        "Formação: Bacharel em Administração Geral"
    ],
    "software_engineer": [
        "Desenvolvedor Full Stack com 5 anos de experiência",
        "Experiência Profissional:",
        "- Desenvolvimento de APIs RESTful usando Python e Node.js",
        "- Trabalho com React, JavaScript e TypeScript",
        "- Experiência com Docker, AWS e DevOps",
        "- Conhecimento em SQL e MongoDB",
        "Formação: Bacharel em Ciência da Computação"
    ],
    
    "marketing": [
        "Especialista em Marketing Digital",
        "Experiência Profissional:",
        "- Gestão de campanhas no Google Ads e Facebook Ads",
        "- Otimização de SEO e análise de métricas",
        "- Estratégias de Social Media e Copywriting",
        "- Experiência com Instagram e LinkedIn Marketing",
        "Formação: Bacharel em Marketing"
    ],
    
    "hr": [
        "Analista de Recursos Humanos Senior",
        "Experiência Profissional:",
        "- Gestão de processos de recrutamento e seleção",
        "- Desenvolvimento de programas de treinamento",
        "- Administração de folha de pagamento e benefícios",
        "- Gestão de cargos e salários",
        "Formação: Bacharel em Administração"
    ],
    
    "support": [
        "Analista de Suporte Técnico",
        "Experiência Profissional:",
        "- Atendimento help desk e service desk",
        "- Resolução de problemas de hardware e software",
        "- Manutenção de infraestrutura de redes",
        "- Gestão de tickets e suporte ao usuário",
        "Formação: Técnico em Informática"
    ]
}

def main():
    input_dir = Path("bau_dos_cvs")
    input_dir.mkdir(exist_ok=True)

    for role, content in cv_contents.items():
        pdf_filename = input_dir / f"cv_{role}.pdf"
        create_pdf(str(pdf_filename), content)
        
        docx_filename = input_dir / f"cv_{role}.docx"
        create_docx(str(docx_filename), content)

if __name__ == "__main__":
    main()